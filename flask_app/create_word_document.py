#!/usr/bin/env python
"""
This script creates a DOCX document from a full text and a dictionary of
{segment_text: feedback} entries. For each segment, it finds its location in
the text, splits the paragraph into runs, and uses low‑level XML manipulation
to add comment range markers. Then it adds a comments.xml part to the DOCX.
"""

import os
import zipfile
import shutil
import tempfile
from datetime import datetime
from lxml import etree

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_comment_to_run(run, comment_id):
    """
    Given a python‑docx run, add the necessary XML elements so that its text is
    “wrapped” by a comment range. Also insert a comment reference run after it.
    """
    r = run._r  # the underlying lxml element for the run

    # Create and insert a commentRangeStart before the run.
    comment_start = OxmlElement('w:commentRangeStart')
    comment_start.set(qn('w:id'), str(comment_id))
    r.addprevious(comment_start)

    # Create and insert a commentRangeEnd after the run.
    comment_end = OxmlElement('w:commentRangeEnd')
    comment_end.set(qn('w:id'), str(comment_id))
    r.addnext(comment_end)

    # Create a new run element for the comment reference.
    parent = r.getparent()
    comment_ref_run = OxmlElement('w:r')
    comment_ref = OxmlElement('w:commentReference')
    comment_ref.set(qn('w:id'), str(comment_id))
    comment_ref_run.append(comment_ref)

    # Insert the comment reference run right after the current run.
    idx = parent.index(r)
    parent.insert(idx + 1, comment_ref_run)


def add_comments_to_docx(docx_path, comments_info):
    """
    Given a DOCX file and a list of comments (tuples of (comment_id, feedback)),
    this function adds a comments.xml part to the DOCX archive and updates the
    [Content_Types].xml and document.xml.rels files accordingly.
    """
    # Create a temporary folder to extract the DOCX (it is a zip archive)
    temp_dir = tempfile.mkdtemp()
    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        zip_ref.extractall(temp_dir)

    # -------------------------------
    # Create the word/comments.xml file
    comments_xml_path = os.path.join(temp_dir, "word", "comments.xml")
    NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    root = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comments", nsmap=NSMAP)
    # For each comment, create a <w:comment> element.
    for cid, feedback in comments_info:
        comment = etree.SubElement(
            root, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment",
            {
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id": str(cid),
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author": "Reviewer",
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date": datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
            }
        )
        p = etree.SubElement(comment, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p")
        r = etree.SubElement(p, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")
        t = etree.SubElement(r, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
        t.text = feedback

    tree = etree.ElementTree(root)
    tree.write(comments_xml_path, xml_declaration=True, encoding="UTF-8", standalone="yes")

    # -------------------------------
    # Update [Content_Types].xml to add an Override for comments.xml if needed.
    content_types_path = os.path.join(temp_dir, "[Content_Types].xml")
    tree = etree.parse(content_types_path)
    root_ct = tree.getroot()
    override_xpath = ".//{http://schemas.openxmlformats.org/package/2006/content-types}Override[@PartName='/word/comments.xml']"
    ns = {'ct': 'http://schemas.openxmlformats.org/package/2006/content-types'}
    if not root_ct.xpath(".//ct:Override[@PartName='/word/comments.xml']", namespaces=ns):
        override = etree.SubElement(
            root_ct, "{http://schemas.openxmlformats.org/package/2006/content-types}Override"
        )
        override.set("PartName", "/word/comments.xml")
        override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml")
        tree.write(content_types_path, xml_declaration=True, encoding="UTF-8", standalone="yes")


    # -------------------------------
    # Update word/_rels/document.xml.rels to add a relationship for comments.
    rels_path = os.path.join(temp_dir, "word", "_rels", "document.xml.rels")
    tree = etree.parse(rels_path)
    root_rels = tree.getroot()
    ns_rel = {"rel": "http://schemas.openxmlformats.org/package/2006/relationships"}
    if not root_rels.xpath(
        ".//rel:Relationship[@Type='http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments']",
        namespaces=ns_rel
    ):
        # Generate a new relationship id not already in use.
        existing_ids = [rel.get("Id") for rel in root_rels.findall("{http://schemas.openxmlformats.org/package/2006/relationships}Relationship")]
        new_id = "rIdComments"
        counter = 1
        while new_id in existing_ids:
            new_id = f"rIdComments{counter}"
            counter += 1
        relationship = etree.SubElement(
            root_rels, "{http://schemas.openxmlformats.org/package/2006/relationships}Relationship"
        )
        relationship.set("Id", new_id)
        relationship.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments")
        relationship.set("Target", "comments.xml")
        tree.write(rels_path, xml_declaration=True, encoding="UTF-8", standalone="yes")

    # -------------------------------
    # Repackage the modified files back into a DOCX.
    new_docx_path = docx_path  # overwrite original
    with zipfile.ZipFile(new_docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_write:
        for foldername, subfolders, filenames in os.walk(temp_dir):
            for filename in filenames:
                file_path = os.path.join(foldername, filename)
                archive_path = os.path.relpath(file_path, temp_dir)
                zip_write.write(file_path, archive_path)

    shutil.rmtree(temp_dir)


def create_docx_with_comments(text, comment_map, output_path):
    # Find the positions of each segment in the full text.
    segments = []
    for segment, feedback in comment_map.items():
        start = text.find(segment)
        if start != -1:
            segments.append((start, start + len(segment), segment, feedback))
        else:
            print(f"Warning: segment '{segment}' not found in text.")
    # Sort segments by their starting index.
    segments.sort(key=lambda x: x[0])

    # Create a new document and add a single paragraph.
    doc = Document()
    p = doc.add_paragraph()
    comments_info = []  # List to store (comment_id, feedback) for later use.
    comment_id = 0

    current = 0
    for start, end, segment, feedback in segments:
        # Add run for text before the commented segment.
        if current < start:
            p.add_run(text[current:start])
        # Add run for the commented segment.
        run = p.add_run(text[start:end])
        # Insert comment markers into this run.
        add_comment_to_run(run, comment_id)
        comments_info.append((comment_id, feedback))
        comment_id += 1
        current = end
    # Add any remaining text after the last segment.
    if current < len(text):
        p.add_run(text[current:])

    # Save the document.
    output_docx = output_path
    doc.save(output_docx)
    print(f"Document saved as {output_docx}")

    # Add the comments part to the DOCX.
    add_comments_to_docx(output_docx, comments_info)
    print("Comments have been added to the document.")


def create_docx_with_comments_with_headings(text_dict, comment_map, output_path):
    """
    Creates a DOCX document where the text is provided as a dictionary {heading: paragraph}
    and inserts comments (using comment_map) into the paragraph text.
    
    :param text_dict: Dictionary where each key is a heading and its value is the paragraph text.
    :param comment_map: Dictionary mapping text segments (str) to comment feedback (str).
    :param output_path: Path where the DOCX file will be saved.
    """
    doc = Document()
    comments_info = []  # List to store (comment_id, feedback) for later use.
    comment_id = 0

    # Process each heading/paragraph pair.
    for heading, paragraph_text in text_dict.items():
        # Add the heading with a desired level (adjust level as needed).
        doc.add_heading(heading, level=1)
        
        # Add a new paragraph.
        p = doc.add_paragraph()
        
        # Find all comment segments within this paragraph.
        segments = []
        for segment, feedback in comment_map.items():
            start = paragraph_text.find(segment)
            if start != -1:
                segments.append((start, start + len(segment), segment, feedback))
            else:
                print(f"Warning: segment '{segment}' not found in paragraph for heading '{heading}'.")
        
        # Sort segments by their starting index.
        segments.sort(key=lambda x: x[0])
        
        # Process the paragraph text, inserting comment markers.
        current = 0
        for start, end, segment, feedback in segments:
            # Add run for text before the commented segment.
            if current < start:
                p.add_run(paragraph_text[current:start])
            # Add run for the commented segment.
            run = p.add_run(paragraph_text[start:end])
            # Insert a comment marker into this run.
            add_comment_to_run(run, comment_id)
            comments_info.append((comment_id, feedback))
            comment_id += 1
            current = end
        
        # Add any remaining text after the last segment.
        if current < len(paragraph_text):
            p.add_run(paragraph_text[current:])

    # Save the document.
    doc.save(output_path)
    print(f"Document saved as {output_path}")

    # Add the comments part to the DOCX.
    add_comments_to_docx(output_path, comments_info)
    print("Comments have been added to the document.")
